# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r'D:\WB_Workflow\开题报告题目分析.docx')

# Print paragraphs 30-80
print('--- Paragraphs 30-80 ---')
for i, p in enumerate(doc.paragraphs):
    if i >= 29 and i < 80:
        if p.text.strip():
            print(f'  [{i}] {p.text.strip()[:120]}')

print('\n--- Paragraphs 80-140 ---')
for i, p in enumerate(doc.paragraphs):
    if i >= 80:
        if p.text.strip():
            print(f'  [{i}] {p.text.strip()[:120]}')

print(f'\n--- Tables ---')
for i, t in enumerate(doc.tables):
    print(f'Table {i+1}: {len(t.rows)} rows x {len(t.columns)} cols')
    # Print header row
    if t.rows:
        header = [cell.text.strip()[:20] for cell in t.rows[0].cells]
        print(f'  Headers: {header}')
