# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r'D:\WB_Workflow\开题报告题目分析.docx')
print(f'Paragraphs: {len(doc.paragraphs)}')
print(f'Tables: {len(doc.tables)}')
print(f'Sections: {len(doc.sections)}')

for i, sec in enumerate(doc.sections):
    header = sec.header
    if header and header.paragraphs:
        for p in header.paragraphs:
            if p.text.strip():
                print(f'Page {i+1} header: {p.text.strip()[:60]}')

key_phrases = ['新质生产力', '人工智能', '四链融合', '导师', '话术', '数据颗粒度', '创新', 'Dagum', '熵值法', 'DID']
for phrase in key_phrases:
    found = sum(1 for p in doc.paragraphs if phrase in p.text)
    print(f'  [{phrase}]: mentioned in {found} paragraphs')

print('\n--- First 30 non-empty paragraphs ---')
count = 0
for p in doc.paragraphs:
    if p.text.strip():
        print(f'  {p.text.strip()[:120]}')
        count += 1
        if count >= 30:
            break
